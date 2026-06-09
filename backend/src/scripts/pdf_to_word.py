import fitz
import sys
import os

def pdf_to_word_gemini(input_path, output_path, gemini_key):
    try:
        import google.generativeai as genai
        import docx
    except ImportError:
        print("Required python packages for Gemini conversion not found, using direct conversion.", file=sys.stderr)
        pdf_to_word_direct(input_path, output_path)
        return

    try:
        genai.configure(api_key=gemini_key)
        with open(input_path, "rb") as f:
            pdf_data = f.read()
            
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = "Extract all text and tabular data from this PDF document page-by-page. Preserve headings, paragraphs, and list structures as closely as possible. Output only the clean extracted markdown/structured text without extra chat conversational wrapper."
        
        response = model.generate_content([
            {
                "mime_type": "application/pdf",
                "data": pdf_data
            },
            prompt
        ])
        text = response.text
    except Exception as e:
        print(f"Gemini API conversion failed: {e}. Falling back to local pdf2docx conversion...", file=sys.stderr)
        pdf_to_word_direct(input_path, output_path)
        return

    # Parse text and write to docx
    try:
        doc = docx.Document()
        lines = text.split("\n")
        for line in lines:
            cleaned = line.strip()
            if not cleaned:
                continue
            if cleaned.startswith("# "):
                doc.add_heading(cleaned[2:], level=1)
            elif cleaned.startswith("## "):
                doc.add_heading(cleaned[3:], level=2)
            elif cleaned.startswith("### "):
                doc.add_heading(cleaned[4:], level=3)
            elif cleaned.startswith("- ") or cleaned.startswith("* "):
                doc.add_paragraph(cleaned[2:], style='List Bullet')
            else:
                doc.add_paragraph(cleaned)
                
        doc.save(output_path)
    except Exception as e:
        print(f"Error saving DOCX from Gemini output: {e}. Falling back to direct conversion...", file=sys.stderr)
        pdf_to_word_direct(input_path, output_path)

def pdf_to_word_direct(input_path, output_path):
    try:
        from pdf2docx import Converter
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
    except Exception as e:
        print(f"Direct PDF to DOCX conversion failed: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word.py <input_path> <output_path> [gemini_api_key]", file=sys.stderr)
        sys.exit(1)
        
    in_p = sys.argv[1]
    out_p = sys.argv[2]
    key = sys.argv[3] if len(sys.argv) > 3 and sys.argv[3].strip() and sys.argv[3] != "undefined" else ""
    
    if key:
        pdf_to_word_gemini(in_p, out_p, key)
    else:
        pdf_to_word_direct(in_p, out_p)
