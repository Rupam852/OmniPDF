import fitz
import sys
import docx

def word_to_pdf(input_path, output_path):
    try:
        doc = docx.Document(input_path)
    except Exception as e:
        print(f"Error loading Word document: {e}", file=sys.stderr)
        sys.exit(1)
        
    html_content = "<html><body style='font-family: Arial, sans-serif; padding: 40px; line-height: 1.6; color: #1e293b;'>"
    
    # Process elements sequentially
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            html_content += "<br/>"
            continue
            
        style_name = p.style.name.lower() if p.style else ""
        
        if "heading 1" in style_name:
            html_content += f"<h1 style='color: #1e3a8a; margin-top: 20px; font-size: 22px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;'>{text}</h1>"
        elif "heading 2" in style_name:
            html_content += f"<h2 style='color: #2563eb; margin-top: 16px; font-size: 18px;'>{text}</h2>"
        elif "heading 3" in style_name:
            html_content += f"<h3 style='color: #1d4ed8; margin-top: 12px; font-size: 15px;'>{text}</h3>"
        elif "list" in style_name:
            html_content += f"<ul style='margin: 4px 0;'><li style='font-size: 11px;'>{text}</li></ul>"
        else:
            p_html = ""
            for run in p.runs:
                r_text = run.text
                if not r_text:
                    continue
                if run.bold:
                    r_text = f"<b>{r_text}</b>"
                if run.italic:
                    r_text = f"<i>{r_text}</i>"
                p_html += r_text
            html_content += f"<p style='font-size: 11px; margin: 8px 0;'>{p_html}</p>"

    # Process tables
    for table in doc.tables:
        html_content += "<table style='border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 10px;'>"
        for row in table.rows:
            html_content += "<tr>"
            for cell in row.cells:
                # Recursively process cell paragraphs
                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                html_content += f"<td style='border: 1px solid #cbd5e1; padding: 6px; min-width: 50px;'>{cell_text}</td>"
            html_content += "</tr>"
        html_content += "</table>"

    html_content += "</body></html>"
    
    try:
        fitz_doc = fitz.open(stream=html_content.encode('utf-8'), filetype='html')
        pdf_bytes = fitz_doc.convert_to_pdf()
        with open(output_path, "wb") as f:
            f.write(pdf_bytes)
        fitz_doc.close()
    except Exception as e:
        print(f"Error generating PDF: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python word_to_pdf.py <input_path> <output_path>", file=sys.stderr)
        sys.exit(1)
    word_to_pdf(sys.argv[1], sys.argv[2])
